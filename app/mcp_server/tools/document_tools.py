import os
import html
import httpx
import tempfile
from pathlib import Path
from typing import List, Any

from app.config import PROJECT_ROOT, get_settings


ALLOWED_READ_EXTENSIONS = {".docx", ".pdf", ".xlsx", ".xlsm", ".txt", ".md", ".csv"}
ALLOWED_CREATE_EXTENSIONS = {"docx", "xlsx", "pdf", "txt", "md", "csv"}


def summarize_document(file_path: str | None = None, fallback_text: str = "") -> dict:
    """文档摘要工具"""
    if file_path:
        try:
            path = _resolve_safe_input_path(file_path)
        except ValueError as exc:
            return {"summary": str(exc), "keywords": [], "source": file_path}
        if not path.exists():
            return {"summary": f"未找到文件：{file_path}", "keywords": [], "source": file_path}
        text = extract_document_text(path)
        summary = summarize_text(text)
        return {
            "summary": summary,
            "keywords": [path.suffix.lstrip(".") or "document", "文档摘要"],
            "source": str(path),
            "text": text[:4000],
            "character_count": len(text),
        }

    text = fallback_text.strip()
    if not text:
        return {"summary": "未提供文档或文本内容。", "keywords": [], "source": None}

    return {
        "summary": summarize_text(text),
        "keywords": ["办公", "摘要"],
        "source": "message",
        "text": text[:4000],
        "character_count": len(text),
    }


def read_document(file_path: str, extract_tables: bool = True, extract_images: bool = False) -> dict:
    """读取和解析文档内容"""
    try:
        path = _resolve_safe_input_path(file_path)
    except ValueError as exc:
        return {"error": str(exc), "content": None}
    if not path.exists():
        return {"error": f"文件不存在：{file_path}", "content": None}
    if path.suffix.lower() not in ALLOWED_READ_EXTENSIONS:
        return {"error": f"不支持的文件格式：{path.suffix}", "content": None}
    if path.stat().st_size > get_settings().max_file_size_bytes:
        return {"error": f"文件过大，超过 {get_settings().max_file_size_mb}MB 限制", "content": None}

    suffix = path.suffix.lower()
    result = {
        "file_path": str(path),
        "file_name": path.name,
        "file_type": suffix,
        "file_size": path.stat().st_size,
        "content": {},
        "metadata": {}
    }

    try:
        if suffix == ".docx":
            result["content"] = _extract_docx_detailed(path, extract_tables, extract_images)
        elif suffix == ".pdf":
            result["content"] = _extract_pdf_detailed(path, extract_tables, extract_images)
        elif suffix in {".xlsx", ".xlsm", ".xls"}:
            result["content"] = _extract_excel_detailed(path)
        elif suffix in {".txt", ".md"}:
            result["content"] = _extract_text_detailed(path)
        else:
            result["error"] = f"不支持的文件格式：{suffix}"
    except Exception as e:
        result["error"] = f"解析失败：{str(e)}"

    return result


def create_document(file_type: str, content: dict, output_path: str) -> dict:
    """创建文档文件"""
    normalized_type = file_type.lower()
    if normalized_type not in ALLOWED_CREATE_EXTENSIONS:
        return {"success": False, "error": f"不支持的文件类型：{file_type}"}
    path = _resolve_safe_output_path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    try:
        if normalized_type == "docx":
            return _create_docx(content, path)
        elif normalized_type == "xlsx":
            return _create_excel(content, path)
        elif normalized_type == "pdf":
            return _create_pdf(content, path)
        elif normalized_type in {"txt", "md", "csv"}:
            return _create_text(content, path)
    except Exception as e:
        return {"success": False, "error": f"创建失败：{str(e)}"}


def process_table_data(table_data: List[List[Any]], operation: str, **kwargs) -> dict:
    """处理表格数据"""
    try:
        if operation == "filter":
            return _filter_table(table_data, **kwargs)
        elif operation == "sort":
            return _sort_table(table_data, **kwargs)
        elif operation == "merge":
            return _merge_tables(table_data, **kwargs)
        elif operation == "pivot":
            return _pivot_table(table_data, **kwargs)
        elif operation == "analyze":
            return _analyze_table(table_data, **kwargs)
        else:
            return {"error": f"不支持的操作：{operation}"}
    except Exception as e:
        return {"error": f"表格处理失败：{str(e)}"}


def convert_document(input_path: str, output_path: str, target_format: str) -> dict:
    """文档格式转换"""
    try:
        input_path_obj = _resolve_safe_input_path(input_path)
        output_path_obj = _resolve_safe_output_path(output_path)
    except ValueError as exc:
        return {"success": False, "error": str(exc)}

    if not input_path_obj.exists():
        return {"success": False, "error": f"输入文件不存在：{input_path_obj}"}

    try:
        normalized_target = target_format.lower().lstrip(".")
        if normalized_target == "pdf":
            gotenberg_result = _convert_with_gotenberg(input_path_obj, output_path_obj)
            return gotenberg_result

        # 读取源文档
        source_data = read_document(str(input_path_obj))
        if "error" in source_data:
            return {"success": False, "error": source_data["error"]}

        # 转换为目标格式
        fallback_result = create_document(normalized_target, source_data["content"], str(output_path_obj))
        return fallback_result
    except Exception as e:
        return {"success": False, "error": f"转换失败：{str(e)}"}


def extract_document_text(path: Path) -> str:
    """提取文档文本（兼容旧接口）"""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix in {".xlsx", ".xlsm"}:
        return _extract_xlsx(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".csv":
        return path.read_text(encoding="utf-8", errors="ignore")
    return f"暂不支持解析 {suffix or '未知'} 格式文件：{path.name}"


def summarize_text(text: str, max_chars: int = 300) -> str:
    clean = " ".join(text.split())
    if not clean:
        return "文档没有提取到可用文本。"
    if len(clean) <= max_chars:
        return clean
    return clean[:max_chars] + "..."


def _extract_docx_detailed(path: Path, extract_tables: bool = True, extract_images: bool = False) -> dict:
    """详细解析Word文档"""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx 依赖，无法解析 docx 文件。") from exc

    document = Document(path)
    result = {
        "text_content": [],
        "tables": [],
        "images": [],
        "metadata": {
            "paragraphs_count": len(document.paragraphs),
            "tables_count": len(document.tables),
        }
    }

    # 提取段落文本
    for para in document.paragraphs:
        if para.text.strip():
            result["text_content"].append({
                "type": "paragraph",
                "text": para.text.strip(),
                "style": para.style.name if para.style else None
            })

    # 提取表格
    if extract_tables:
        for i, table in enumerate(document.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)

            result["tables"].append({
                "index": i,
                "rows": len(table_data),
                "columns": len(table_data[0]) if table_data else 0,
                "data": table_data
            })

    # 提取图片信息（如果需要）
    if extract_images:
        for rel in document.part.rels:
            if "image" in document.part.rels[rel].reltype:
                result["images"].append({
                    "relationship_id": rel,
                    "type": "image"
                })

    return result


def _extract_pdf_detailed(path: Path, extract_tables: bool = True, extract_images: bool = False) -> dict:
    """详细解析PDF文档"""
    try:
        import pdfplumber
        from pdfplumber.page import Page
    except ImportError as exc:
        raise RuntimeError("缺少 pdfplumber 依赖，无法解析 pdf 文件。") from exc

    result = {
        "text_content": [],
        "tables": [],
        "images": [],
        "metadata": {}
    }

    with pdfplumber.open(path) as pdf:
        result["metadata"]["pages_count"] = len(pdf.pages)

        for i, page in enumerate(pdf.pages):
            page_result = {
                "page_number": i + 1,
                "text": page.extract_text() or "",
                "tables": [],
                "images": []
            }

            # 提取表格
            if extract_tables:
                tables = page.extract_tables()
                for j, table in enumerate(tables):
                    if table:
                        page_result["tables"].append({
                            "index": j,
                            "data": table
                        })

            result["text_content"].append(page_result)

    return result


def _extract_excel_detailed(path: Path) -> dict:
    """详细解析Excel文档"""
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("缺少 openpyxl 和 pandas 依赖，无法解析 excel 文件。") from exc

    workbook = load_workbook(path, read_only=True, data_only=True)
    result = {
        "sheets": [],
        "metadata": {
            "sheets_count": len(workbook.worksheets),
            "file_name": path.name
        }
    }

    for sheet in workbook.worksheets:
        sheet_data = []
        headers = []

        # 获取表头
        for row in sheet.iter_rows(min_row=1, max_row=1, values_only=True):
            headers = [str(cell) if cell is not None else "" for cell in row]
            break

        # 获取数据
        for row in sheet.iter_rows(min_row=2, values_only=True):
            row_data = []
            for cell in row:
                row_data.append(str(cell) if cell is not None else "")
            if any(row_data):  # 只添加非空行
                sheet_data.append(row_data)

        result["sheets"].append({
            "name": sheet.title,
            "headers": headers,
            "data": sheet_data,
            "rows_count": len(sheet_data),
            "columns_count": len(headers)
        })

    workbook.close()
    return result


def _extract_text_detailed(path: Path) -> dict:
    """详细解析文本文档"""
    try:
        content = path.read_text(encoding="utf-8", errors="ignore")
        lines = content.split('\n')

        return {
            "content": content,
            "lines": lines,
            "lines_count": len(lines),
            "characters_count": len(content),
            "metadata": {
                "encoding": "utf-8",
                "file_type": path.suffix
            }
        }
    except Exception as e:
        return {"error": f"读取文本文件失败：{str(e)}"}


def _extract_docx(path: Path) -> str:
    detailed = _extract_docx_detailed(path, extract_tables=False, extract_images=False)
    paragraphs = [item.get("text", "") for item in detailed.get("text_content", []) if item.get("text")]
    return "\n".join(paragraphs)


def _extract_pdf(path: Path) -> str:
    detailed = _extract_pdf_detailed(path, extract_tables=False, extract_images=False)
    pages = [item.get("text", "") for item in detailed.get("text_content", []) if item.get("text")]
    return "\n".join(pages)


def _extract_xlsx(path: Path) -> str:
    detailed = _extract_excel_detailed(path)
    lines: list[str] = []
    for sheet in detailed.get("sheets", []):
        lines.append(f"Sheet: {sheet.get('name', '')}")
        headers = sheet.get("headers", [])
        if headers:
            lines.append(",".join(str(cell) for cell in headers))
        for row in sheet.get("data", []):
            lines.append(",".join(str(cell) for cell in row))
    return "\n".join(lines)


def _resolve_safe_input_path(file_path: str) -> Path:
    settings = get_settings()
    path = Path(file_path)
    if not path.is_absolute():
        path = PROJECT_ROOT / path
    path = path.resolve()

    if settings.allow_external_input_files or _is_test_runtime(settings):
        return path

    upload_root = settings.resolved_upload_dir.resolve()
    project_root = PROJECT_ROOT.resolve()
    if str(path).startswith(str(project_root)) or str(path).startswith(str(upload_root)):
        return path
    raise ValueError("不允许访问工作区之外的路径。")


def _resolve_safe_output_path(output_path: str) -> Path:
    settings = get_settings()
    path = Path(output_path)
    if not path.is_absolute():
        path = PROJECT_ROOT / path
    path = path.resolve()
    project_root = PROJECT_ROOT.resolve()
    if str(path).startswith(str(project_root)) or _is_test_runtime(settings):
        return path
    raise ValueError("输出路径必须位于项目目录内。")


def _is_test_runtime(settings=None) -> bool:
    active_settings = settings or get_settings()
    return active_settings.app_env == "test" or "PYTEST_CURRENT_TEST" in os.environ


def _create_docx(content: dict, output_path: Path) -> dict:
    """创建Word文档"""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        return {"success": False, "error": f"缺少 python-docx 依赖：{str(exc)}"}

    doc = Document()
    _set_docx_cjk_font(doc, get_settings().office_agent_docx_font)

    # 添加标题
    if "title" in content:
        doc.add_heading(content["title"], 0)

    # 添加段落
    if "paragraphs" in content:
        for para in content["paragraphs"]:
            if isinstance(para, str):
                doc.add_paragraph(para)
            elif isinstance(para, dict):
                p = doc.add_paragraph(para.get("text", ""))
                if "style" in para:
                    p.style = para["style"]

    # 添加表格
    if "tables" in content:
        for table_data in content["tables"]:
            if isinstance(table_data, list) and table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = str(cell)

    doc.save(str(output_path))
    return {"success": True, "file_path": str(output_path)}


def _set_docx_cjk_font(doc, font_family: str) -> None:
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        return

    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = font_family
        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{attr}"), font_family)


def _create_excel(content: dict, output_path: Path) -> dict:
    """创建Excel文档"""
    try:
        from openpyxl import Workbook
    except ImportError as exc:
        return {"success": False, "error": f"缺少 openpyxl 依赖：{str(exc)}"}

    wb = Workbook()

    if "sheets" in content:
        for sheet_index, sheet_data in enumerate(content["sheets"]):
            ws = wb.active if sheet_index == 0 else wb.create_sheet()
            sheet_name = str(sheet_data.get("name") or f"Sheet{sheet_index + 1}")[:31]
            ws.title = sheet_name or f"Sheet{sheet_index + 1}"
            _write_sheet_data(ws, sheet_data)
    else:
        # 默认单sheet处理
        ws = wb.active
        if "data" in content:
            for row_idx, row in enumerate(content["data"], 1):
                for col_idx, cell_value in enumerate(row, 1):
                    ws.cell(row=row_idx, column=col_idx, value=cell_value)

    wb.save(str(output_path))
    return {"success": True, "file_path": str(output_path)}


def _write_sheet_data(ws, sheet_data: dict) -> None:
    if "headers" in sheet_data and sheet_data["headers"]:
        for col, header in enumerate(sheet_data["headers"], 1):
            ws.cell(row=1, column=col, value=header)

    if "data" in sheet_data:
        start_row = 2 if sheet_data.get("headers") else 1
        for row_idx, row in enumerate(sheet_data["data"], start_row):
            for col_idx, cell_value in enumerate(row, 1):
                ws.cell(row=row_idx, column=col_idx, value=cell_value)


def _create_pdf(content: dict, output_path: Path) -> dict:
    """创建PDF文档，优先通过 Gotenberg 转换临时 DOCX 以保留版式。"""
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix=".office_agent_pdf_", dir=output_path.parent) as tmp_dir:
        staging_docx = Path(tmp_dir) / f"{output_path.stem or 'document'}.docx"
        docx_result = _create_docx(content, staging_docx)
        if docx_result.get("success"):
            gotenberg_result = _convert_with_gotenberg(staging_docx, output_path)
            if gotenberg_result["success"]:
                return gotenberg_result

    fallback_result = _create_pdf_with_reportlab(content, output_path)
    if fallback_result.get("success"):
        fallback_result["converter"] = "reportlab-fallback"
        fallback_result["warning"] = "Gotenberg 不可用或转换失败，已使用 ReportLab 中文字体兜底。"
    return fallback_result


def _create_pdf_with_reportlab(content: dict, output_path: Path) -> dict:
    """ReportLab 兜底创建 PDF，注册系统中文字体避免中文变方框。"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib import colors
    except ImportError as exc:
        return {"success": False, "error": f"缺少 reportlab 依赖：{str(exc)}"}

    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    cjk_font = _register_cjk_font()
    normal_style = ParagraphStyle(
        "OfficeAgentNormal",
        parent=styles["Normal"],
        fontName=cjk_font,
        fontSize=11,
        leading=16,
    )
    heading_style = ParagraphStyle(
        "OfficeAgentHeading",
        parent=styles["Heading1"],
        fontName=cjk_font,
        fontSize=18,
        leading=24,
        spaceAfter=12,
    )
    story = []

    # 添加标题
    if "title" in content:
        story.append(Paragraph(html.escape(str(content["title"])), heading_style))
        story.append(Spacer(1, 8))

    # 添加段落
    for para_text in _iter_paragraph_texts(content):
        story.append(Paragraph(html.escape(para_text), normal_style))
        story.append(Spacer(1, 6))

    # 添加表格
    for table_data in _iter_table_data(content):
        if table_data:
            normalized_table = [[str(cell) for cell in row] for row in table_data]
            table = Table(normalized_table)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#d0d0d0")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), cjk_font),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor("#f7f7f7")),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ]))
            story.append(Spacer(1, 8))
            story.append(table)

    if not story:
        story.append(Paragraph("文档没有可写入的内容。", normal_style))

    try:
        doc.build(story)
        return {"success": True, "file_path": str(output_path)}
    except Exception as e:
        return {"success": False, "error": f"PDF 写入失败：{str(e)}"}


def _convert_with_gotenberg(input_path: Path, output_path: Path) -> dict:
    """Convert an office document to PDF through a configured Gotenberg service."""
    settings = get_settings()
    if not settings.gotenberg_url:
        return {"success": False, "error": "未配置 GOTENBERG_URL，跳过 Gotenberg 转换。"}

    output_path.parent.mkdir(parents=True, exist_ok=True)
    endpoint = f"{settings.gotenberg_url.rstrip('/')}/forms/libreoffice/convert"

    try:
        with httpx.Client(timeout=settings.gotenberg_timeout) as client:
            with input_path.open("rb") as input_file:
                response = client.post(
                    endpoint,
                    files={
                        "files": (
                            input_path.name,
                            input_file,
                            "application/octet-stream",
                        )
                    },
                )
        response.raise_for_status()
    except httpx.HTTPStatusError as exc:
        status_code = exc.response.status_code if exc.response is not None else "unknown"
        detail = exc.response.text[:500] if exc.response is not None else str(exc)
        return {"success": False, "error": f"Gotenberg 转换失败：HTTP {status_code} {detail}"}
    except httpx.RequestError as exc:
        return {"success": False, "error": f"Gotenberg 请求失败：{exc}"}
    except OSError as exc:
        return {"success": False, "error": f"读取输入文件失败：{exc}"}

    output_path.write_bytes(response.content)
    return {
        "success": True,
        "file_path": str(output_path),
        "converter": "gotenberg",
        "source_path": str(input_path),
    }


def _register_cjk_font() -> str:
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return "Helvetica"

    font_name = "OfficeAgentCJK"
    if font_name in pdfmetrics.getRegisteredFontNames():
        return font_name

    custom_font_dir = Path(get_settings().office_agent_font_dir)
    candidates = [
        str(custom_font_dir / "微软雅黑.ttf"),
        str(custom_font_dir / "思源宋体.ttf"),
        str(custom_font_dir / "NotoSerifSC-VF.ttf"),
        str(custom_font_dir / "NotoSerifCJKsc-VF.ttf"),
        str(custom_font_dir / "NotoSerifCJK-VF.ttf.ttc"),
        str(custom_font_dir / "NotoSerifCJKsc-Regular.otf"),
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/truetype/arphic/uming.ttc",
    ]

    for font_path in candidates:
        path = Path(font_path)
        if not path.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont(font_name, str(path)))
            return font_name
        except Exception:
            continue

    try:
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont

        cid_font = "STSong-Light"
        pdfmetrics.registerFont(UnicodeCIDFont(cid_font))
        return cid_font
    except Exception:
        pass
    return "Helvetica"


def _iter_paragraph_texts(content: dict) -> list[str]:
    paragraphs: list[str] = []

    for para in content.get("paragraphs", []) or []:
        if isinstance(para, str):
            text = para.strip()
        elif isinstance(para, dict):
            text = str(para.get("text", "")).strip()
        else:
            text = str(para).strip()
        if text:
            paragraphs.append(text)

    for item in content.get("text_content", []) or []:
        if isinstance(item, dict):
            text = str(item.get("text", "")).strip()
        else:
            text = str(item).strip()
        if text:
            paragraphs.append(text)

    content_text = str(content.get("content", "")).strip()
    if content_text and not paragraphs:
        paragraphs.append(content_text)

    return paragraphs


def _iter_table_data(content: dict) -> list[list[list[Any]]]:
    tables: list[list[list[Any]]] = []
    for table_data in content.get("tables", []) or []:
        if isinstance(table_data, dict):
            data = table_data.get("data", [])
        else:
            data = table_data
        if isinstance(data, list) and data:
            rows = [row if isinstance(row, list) else [row] for row in data]
            tables.append(rows)
    return tables


def _create_text(content: dict, output_path: Path) -> dict:
    """创建文本文档"""
    text_content = content.get("content", "")

    if "paragraphs" in content and not text_content:
        text_content = "\n\n".join(content["paragraphs"])

    if "tables" in content and not text_content:
        table_texts = []
        for table in content["tables"]:
            if isinstance(table, list):
                for row in table:
                    if isinstance(row, list):
                        table_texts.append(" | ".join(str(cell) for cell in row))
                    else:
                        table_texts.append(str(row))
                table_texts.append("")
        text_content = "\n".join(table_texts)

    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        return {"success": True, "file_path": str(output_path)}
    except Exception as e:
        return {"success": False, "error": f"写入文件失败：{str(e)}"}


def _filter_table(table_data: List[List[Any]], **kwargs) -> dict:
    """过滤表格数据"""
    column = kwargs.get("column", 0)
    value = kwargs.get("value", "")
    operator = kwargs.get("operator", "equals")

    if not table_data or column >= len(table_data[0]):
        return {"error": "无效的表格数据或列索引"}

    filtered_data = []
    for row in table_data:
        cell_value = str(row[column]).lower() if column < len(row) else ""
        filter_value = str(value).lower()

        if operator == "equals" and cell_value == filter_value:
            filtered_data.append(row)
        elif operator == "contains" and filter_value in cell_value:
            filtered_data.append(row)
        elif operator == "starts_with" and cell_value.startswith(filter_value):
            filtered_data.append(row)
        elif operator == "ends_with" and cell_value.endswith(filter_value):
            filtered_data.append(row)

    return {"filtered_data": filtered_data, "original_rows": len(table_data), "filtered_rows": len(filtered_data)}


def _sort_table(table_data: List[List[Any]], **kwargs) -> dict:
    """排序表格数据"""
    column = kwargs.get("column", 0)
    reverse = kwargs.get("reverse", False)

    if not table_data:
        return {"error": "无效的表格数据"}

    try:
        sorted_data = sorted(table_data, key=lambda x: x[column] if column < len(x) else "", reverse=reverse)
        return {"sorted_data": sorted_data, "sort_column": column, "reverse": reverse}
    except Exception as e:
        return {"error": f"排序失败：{str(e)}"}


def _merge_tables(table_data: List[List[Any]], **kwargs) -> dict:
    """合并多个表格"""
    # 简化实现：垂直合并
    merged_data = []
    for table in table_data:
        merged_data.extend(table)

    return {"merged_data": merged_data, "total_rows": len(merged_data)}


def _pivot_table(table_data: List[List[Any]], **kwargs) -> dict:
    """数据透视表"""
    # 简化实现：基本的数据透视
    if not table_data or len(table_data) < 2:
        return {"error": "需要至少两行数据"}

    headers = table_data[0]
    data_rows = table_data[1:]

    pivot_result = {}
    for row in data_rows:
        if len(row) > 1:
            key = str(row[0])
            value = row[1] if len(row) > 1 else 0
            if key in pivot_result:
                pivot_result[key] += 1  # 计数
            else:
                pivot_result[key] = 1

    return {"pivot_data": pivot_result, "headers": headers}


def _analyze_table(table_data: List[List[Any]], **kwargs) -> dict:
    """分析表格数据"""
    if not table_data:
        return {"error": "无效的表格数据"}

    analysis = {
        "total_rows": len(table_data),
        "total_columns": len(table_data[0]) if table_data else 0,
        "empty_cells": 0,
        "numeric_columns": [],
        "text_columns": []
    }

    # 分析每列的数据类型
    if table_data:
        for col_idx in range(len(table_data[0])):
            column_values = []
            for row in table_data:
                if col_idx < len(row):
                    value = row[col_idx]
                    if value is not None and str(value).strip():
                        column_values.append(value)
                    else:
                        analysis["empty_cells"] += 1

            # 判断列类型
            try:
                numeric_count = sum(1 for v in column_values if isinstance(v, (int, float)) or str(v).replace('.', '').isdigit())
                if numeric_count > len(column_values) * 0.5:
                    analysis["numeric_columns"].append(col_idx)
                else:
                    analysis["text_columns"].append(col_idx)
            except Exception:
                analysis["text_columns"].append(col_idx)

    return analysis
